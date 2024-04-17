import diagrams
from docx import Document
from docx.shared import Inches
import functions as loading
import pandas as pd


def texts(year, result_table, subject_table):
    temp = result_table[result_table['subject_form_id'].astype(str).str.contains(str(year))]
    temp_prev = result_table[result_table['subject_form_id'].astype(str).str.contains(str(year - 1))]
    count_of_exams = temp['subject_form_id'].nunique()
    count_of_students = temp['student_id'].nunique()
    count_of_students_prev = temp_prev['student_id'].nunique()

    subjects_ids = [str(i).replace(str(year), '') for i in temp.subject_form_id.unique()]
    subjects = list(subject_table[subject_table['subject_id'].isin([int(i) for i in subjects_ids])].subject_name)

    distinction_by_year = count_of_students - temp_prev.student_id.nunique()
    if distinction_by_year < 0:
        distinction_by_year = f'{abs(distinction_by_year)} человек меньше'
    elif distinction_by_year > 0:
        distinction_by_year = f'{abs(distinction_by_year)} человек больше'

    fail_count = temp.query('result_5 == 2').student_id.nunique()
    fail_count_prev = temp_prev.query('result_5 == 2').student_id.nunique()
    fail_percentage = round(fail_count / count_of_students * 100, 2)
    fail_percentage_prev = round(fail_count_prev / count_of_students * 100, 2)

    russian_fail_count = temp.query(f'result_5 == 2 and subject_form_id == {year}01').student_id.nunique()
    russian_fail_count_prev = temp_prev.query(f'result_5 == 2 and subject_form_id == {year - 1}01').student_id.nunique()

    math_fail_count = temp.query(f'result_5 == 2 and subject_form_id == {year}02').student_id.nunique()
    math_fail_count_prev = temp_prev.query(f'result_5 == 2 and subject_form_id == {year - 1}02').student_id.nunique()

    hundred_graders = temp.query('score_100 == 100')
    hundred_grade_cnt = hundred_graders.student_id.nunique()
    hundred_grade_cnt_prev = temp_prev.query('score_100 == 100').student_id.nunique()

    hundred_graders_value = [hundred_graders.query(f'subject_form_id == {year}{i}').student_id.count() for i in
                             subjects_ids]
    hundred_graders_list = dict(zip(subjects, hundred_graders_value))
    for i in list(hundred_graders_list.keys()):
        if hundred_graders_list[i] == 0:
            hundred_graders_list.pop(i)

    if fail_percentage - fail_percentage_prev < 0:
        subj_grade = 'ненамного хуже'
    else:
        subj_grade = 'существенно лучше'

    students_on_three = len(temp.student_id.value_counts().loc[lambda x: x >= 3])
    students_on_three_prev = len(temp_prev.student_id.value_counts().loc[lambda x: x >= 3])
    print(students_on_three)
    print(students_on_three_prev)
    if students_on_three - students_on_three_prev > 0:
        subj_grade_of_stud_cnt_on_three = 'Увеличилась'
    else:
        subj_grade_of_stud_cnt_on_three = 'Упала'

    f1 = f'''
    В {year} году ЕГЭ в Тюменской области на стадии итоговой аттестации проводился по {count_of_exams} общеобразовательным предметам: {', '.join(i for i in subjects)}
    В ЕГЭ в {year} году в Тюменской области участвовало {count_of_students} человек, что на {distinction_by_year} чем в прошлом году.
    '''  # 703 человека больше/меньше

    f2 = f'''
    Результаты ЕГЭ {year} года оказались {subj_grade}, чем в прошлом году. Процент не справившихся с экзаменационными заданиями
    составил {fail_percentage}% от общего числа участников средних образовательных учреждений, что на {abs(fail_percentage - fail_percentage_prev)}% больше чем в прошлом году.
    В {year} году русский язык в форме ЕГЭ не сдали {russian_fail_count} учащихся средних образовательных учреждений ({russian_fail_count_prev} в {year - 1} году).
    По математике достаточного для зачета баллов не набрали {math_fail_count} учеников средних общеобразовательных учреждений ({math_fail_count_prev} в {year - 1} году).
    Зато количество 100 - балльных работ возросло по сравнению с {year - 1} годом ({hundred_grade_cnt_prev} человек). Так в {year} году в Тюменской области
    {hundred_grade_cnt} учащихся получили максимальный балл – 100 баллов: {", ".join("{}: {}".format(*i) for i in hundred_graders_list.items())}.

    '''  # subj_grade существенно лучше/ненамного хуже, 100_graders - 10 по русскому, 2 по математике etc.

    f3 = f'''
    {subj_grade_of_stud_cnt_on_three} численность выпускников средних общеобразовательных учреждений выбравших три и более предмета.
    '''  # увеличилась/упала

    return (f1, f2, f3)


def main():
    year = int(input('Введите год в формате YYYY: '))
    result_table = pd.DataFrame(loading.print_result_table(),
                                columns=['result_id', 'student_id', 'subject_form_id', 'primary_score', 'accuracy',
                                         'score_100', 'result_5'])
    subject_table = pd.DataFrame(loading.print_subject_table(), columns=['subject_id', 'subject_name'])

    texts_data = texts(year, result_table, subject_table)

    fig2 = diagrams.show_participant_count([year - 1, year], functions.get_subject_list(), functions.get_school_CODE_list())
    fig1 = diagrams.average_subject_result_show([year - 1, year], functions.get_subject_list(), functions.get_school_CODE_list())
    image_fig1 = fig1.to_image(format="png")
    image_fig2 = fig2.to_image(format="png")

    text_fig = ""
    trace = fig1.data[1]
    x_values = trace.x
    y_values = trace.y
    data_dict = {"X data": x_values, "Y data": y_values}
    for subject, score in zip(data_dict['X data'], data_dict['Y data']):
        text_fig += f"{subject} имеет средний балл - {score}\n"

    with open('average_subject_result.png', 'wb') as img_file:
        img_file.write(image_fig1)
    with open('participant_count.png', 'wb') as img_file:
        img_file.write(image_fig2)

    document = Document()
    document.add_heading('Итоги Единого государственного экзамена в Тюменской области', level=1)
    document.add_paragraph(f'{texts_data[0]}')
    document.add_picture('participant_count.png', width=Inches(6))
    document.add_picture('average_subject_result.png', width=Inches(6))
    document.add_paragraph(f'{text_fig}')
    document.add_paragraph(f'{texts_data[1]}')
    document.save('test.docx')


main()